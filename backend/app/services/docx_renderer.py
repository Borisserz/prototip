"""Markdown -> DOCX рендерер (Phase 2: генерация отчётов в Word).

Принимает размеченный Markdown (как правило, от LLM) и собирает .docx через
python-docx. Поддерживает:
  - заголовки (#, ##, ###, ####)
  - абзацы с inline-разметкой (**bold**, *italic*, `code`, [текст](url))
  - маркированные списки (-, *, +) и нумерованные (1.)
  - таблицы в стиле GitHub (| a | b | / |---|---|)
  - изображения ![alt](src), где src — локальный путь, http(s)-URL или
    объект в MinIO (резолвится через core.storage.download_to_path)
  - горизонтальные разделители (---)
  - цитаты (>)

Картинки графиков "вытягиваются из MinIO": если src не существует локально и
не является http(s)-URL, он трактуется как ключ объекта в MinIO и скачивается
во временный файл.

Модуль не падает на ошибках вставки картинок — пропускает их с placeholder-текстом.
"""

from __future__ import annotations

import contextlib
import logging
import re
import tempfile
from collections.abc import Callable
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from core import storage

logger = logging.getLogger("DocxRenderer")

# Максимальная ширина картинки в документе (A4 - поля)
DEFAULT_IMG_WIDTH_IN = 6.0

# ─── inline-разметка ────────────────────────────────────────────────────────
_INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`|\[.+?\]\(.+?\))"
)
_LINK_RE = re.compile(r"\[(?P<text>.+?)\]\((?P<url>.+?)\)")
_IMG_RE = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<src>.+?)\)\s*$")


def _default_image_resolver(src: str) -> str | None:
    """Возвращает локальный путь к картинке (скачивает при необходимости)."""
    src = src.strip()
    # 1. Локальный файл
    p = Path(src)
    if p.exists():
        return str(p)
    # 2. http(s) — скачиваем
    if src.startswith("http://") or src.startswith("https://"):
        try:
            import requests

            r = requests.get(src, timeout=20)
            r.raise_for_status()
            suffix = Path(src.split("?")[0]).suffix or ".png"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)  # noqa: SIM115
            tmp.write(r.content)
            tmp.close()
            return tmp.name
        except Exception as e:
            logger.warning("DocxRenderer: не удалось скачать %s: %s", src, e)
            return None
    # 3. Объект в MinIO (ключ вида "charts/foo.png" или "foo.png")
    suffix = Path(src).suffix or ".png"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)  # noqa: SIM115
    tmp.close()
    bucket = None
    object_name = src
    # допускаем формат "bucket:object"
    if ":" in src and not src.startswith("/"):
        maybe_bucket, _, rest = src.partition(":")
        if "/" in rest or "." in rest:
            bucket, object_name = maybe_bucket, rest
    got = storage.download_to_path(object_name, tmp.name, bucket=bucket)
    return got


class MarkdownDocxRenderer:
    """Конвертер Markdown -> python-docx Document."""

    def __init__(
        self,
        image_resolver: Callable[[str], str | None] | None = None,
        img_width_in: float = DEFAULT_IMG_WIDTH_IN,
    ) -> None:
        self.resolve_image = image_resolver or _default_image_resolver
        self.img_width_in = img_width_in

    # ── публичный API ───────────────────────────────────────────────────────
    def render(
        self,
        markdown_text: str,
        out_path: str | Path,
        *,
        title: str | None = None,
        subtitle: str | None = None,
    ) -> Path:
        doc = Document()
        self._setup_styles(doc)

        if title:
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        self._render_blocks(doc, markdown_text or "")

        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        logger.info("DocxRenderer: сохранён %s", out)
        return out

    # ── настройка стилей/страницы ────────────────────────────────────────────
    def _setup_styles(self, doc: Document) -> None:
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)
        except Exception:
            pass
        for section in doc.sections:
            # A4
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)

    # ── блочный парсер ────────────────────────────────────────────────────────
    def _render_blocks(self, doc: Document, text: str) -> None:
        lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        i = 0
        n = len(lines)
        while i < n:
            line = lines[i]
            stripped = line.strip()

            # пустая строка
            if not stripped:
                i += 1
                continue

            # картинка (отдельной строкой)
            m_img = _IMG_RE.match(stripped)
            if m_img:
                self._add_image(doc, m_img.group("src"), m_img.group("alt"))
                i += 1
                continue

            # заголовки
            m_h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m_h:
                level = min(len(m_h.group(1)), 4)
                doc.add_heading(m_h.group(2).strip(), level=level)
                i += 1
                continue

            # горизонтальный разделитель
            if re.match(r"^([-*_])\1{2,}$", stripped):
                p = doc.add_paragraph()
                p.add_run("─" * 40).font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
                i += 1
                continue

            # таблица (заголовок + разделитель |---|)
            if stripped.startswith("|") and i + 1 < n and re.match(
                r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1].strip()
            ) and "-" in lines[i + 1]:
                i = self._add_table(doc, lines, i)
                continue

            # цитата
            if stripped.startswith(">"):
                quote = re.sub(r"^>\s?", "", stripped)
                p = doc.add_paragraph(style="Intense Quote" if self._has_style(doc, "Intense Quote") else None)
                self._add_inline(p, quote)
                i += 1
                continue

            # списки (с учётом вложенности по отступам)
            m_ul = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
            m_ol = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
            if m_ul or m_ol:
                m = m_ul or m_ol
                indent = len(m.group(1))
                level = min(indent // 2, 2)
                style = "List Bullet" if m_ul else "List Number"
                if level > 0:
                    style = f"{style} {level + 1}" if self._has_style(doc, f"{style} {level + 1}") else style
                p = doc.add_paragraph(style=style if self._has_style(doc, style) else None)
                if not self._has_style(doc, style):
                    p.add_run(("    " * level) + ("• " if m_ul else "1. "))
                self._add_inline(p, m.group(2))
                i += 1
                continue

            # обычный абзац (склеиваем последовательные не-пустые строки)
            buf = [stripped]
            j = i + 1
            while j < n and lines[j].strip() and not self._is_block_start(lines[j], lines, j):
                buf.append(lines[j].strip())
                j += 1
            p = doc.add_paragraph()
            self._add_inline(p, " ".join(buf))
            i = j

    def _is_block_start(self, line: str, lines: list[str], idx: int) -> bool:
        s = line.strip()
        if re.match(r"^#{1,6}\s+", s):
            return True
        if re.match(r"^([-*_])\1{2,}$", s):
            return True
        if re.match(r"^(\s*)[-*+]\s+", line) or re.match(r"^(\s*)\d+[.)]\s+", line):
            return True
        if s.startswith(">"):
            return True
        if s.startswith("|"):
            return True
        return bool(_IMG_RE.match(s))

    # ── таблицы ───────────────────────────────────────────────────────────────
    def _split_row(self, row: str) -> list[str]:
        row = row.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [c.strip() for c in row.split("|")]

    def _add_table(self, doc: Document, lines: list[str], start: int) -> int:
        header = self._split_row(lines[start])
        i = start + 2  # пропускаем строку-разделитель |---|
        rows: list[list[str]] = []
        n = len(lines)
        while i < n and lines[i].strip().startswith("|"):
            rows.append(self._split_row(lines[i]))
            i += 1

        ncols = len(header)
        table = doc.add_table(rows=1, cols=ncols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            with contextlib.suppress(Exception):
                table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for c, txt in enumerate(header):
            if c < ncols:
                hdr_cells[c].text = ""
                run = hdr_cells[c].paragraphs[0].add_run(txt)
                run.bold = True
        for r in rows:
            cells = table.add_row().cells
            for c in range(ncols):
                val = r[c] if c < len(r) else ""
                cells[c].text = ""
                self._add_inline(cells[c].paragraphs[0], val)
        doc.add_paragraph()  # отступ после таблицы
        return i

    # ── картинки ────────────────────────────────────────────────────────────
    def _add_image(self, doc: Document, src: str, alt: str = "") -> None:
        local = None
        try:
            local = self.resolve_image(src)
        except Exception as e:
            logger.warning("DocxRenderer: resolver упал на %s: %s", src, e)
        if not local or not Path(local).exists():
            p = doc.add_paragraph()
            run = p.add_run(f"[изображение недоступно: {alt or src}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return
        try:
            doc.add_picture(local, width=Inches(self.img_width_in))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(alt)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        except Exception as e:
            logger.warning("DocxRenderer: не удалось вставить %s: %s", local, e)
            doc.add_paragraph(f"[ошибка вставки изображения: {alt or src}]")

    # ── inline-разметка ───────────────────────────────────────────────────────
    def _add_inline(self, paragraph, text: str) -> None:
        pos = 0
        for m in _INLINE_RE.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            token = m.group(0)
            if token.startswith("**") or token.startswith("__"):
                paragraph.add_run(token[2:-2]).bold = True
            elif token.startswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
            elif token.startswith("[") and "](" in token:
                lm = _LINK_RE.match(token)
                if lm:
                    run = paragraph.add_run(lm.group("text"))
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                    run.underline = True
                else:
                    paragraph.add_run(token)
            elif token.startswith("*") or token.startswith("_"):
                paragraph.add_run(token[1:-1]).italic = True
            else:
                paragraph.add_run(token)
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    # ── утилиты ───────────────────────────────────────────────────────────────
    @staticmethod
    def _has_style(doc: Document, name: str) -> bool:
        try:
            _ = doc.styles[name]
            return True
        except Exception:
            return False


def markdown_to_docx(
    markdown_text: str,
    out_path: str | Path,
    *,
    title: str | None = None,
    subtitle: str | None = None,
    image_resolver: Callable[[str], str | None] | None = None,
) -> Path:
    """Удобная функция-обёртка."""
    return MarkdownDocxRenderer(image_resolver=image_resolver).render(
        markdown_text, out_path, title=title, subtitle=subtitle
    )
