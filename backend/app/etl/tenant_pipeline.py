"""— ядро ETL-процесса одного клиента (мультитенант).

Переиспользует готовые строительные блоки из ``scripts/build_tenant.py``
(маппинг типов, чтение Postgres, генерация DDL, индексация семантики) и
``scripts/schema_discovery.py``, оборачивая их в одну параметризованную по
``client_id`` функцию ``run_tenant_etl(...)``.

процесс полностью контейнеро-независим: данные грузятся в УЖЕ существующий
ClickHouse клиента (host/port/db из реестра tenant'ов), отдельные docker-контейнеры
поднимать не нужно — это делает запуск из Airflow надёжным (без docker-in-docker).

Шаги:
  1. EXTRACT  — read-only выборка схемы и данных из Postgres клиента.
  2. LOAD     — DDL + INSERT в ClickHouse клиента (БД tenant_<id>).
  3. SEMANTIC — генерация и индексация семантических векторов колонок
                (понимание БД для LLM) в личную коллекцию клиента.
  4. DOCS     — (опционально) загрузка документации клиента в его RAG-коллекцию.

Модуль устойчив к отсутствию embeddings-модели и частичным сбоям: каждый шаг
логируется, итог возвращается структурой ``EtlResult``.
"""

from __future__ import annotations

import logging
import os
import sys
import uuid
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

# корень backend/ в PYTHONPATH (чтобы импортировать scripts.* / app.* / core.*)
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

logger = logging.getLogger("etl.tenant_pipeline")

# переиспользуем строительные блоки из CLI-скрипта build_tenant.py
from scripts.build_tenant import (  # noqa: E402
    build_ch_ddl,
    build_semantics,
    connect_pg,
    fetch_rows,
    get_columns,
    index_semantics,
    list_tables,
)

ProgressCb = Callable[[str, str], None]  # (step, message)


def _utcnow() -> str:
    return datetime.now(UTC).isoformat()


@dataclass
class EtlResult:
    client_id: str
    status: str = "running"  # running | success | failed
    tables: list[str] = field(default_factory=list)
    rows_loaded: dict[str, int] = field(default_factory=dict)
    total_rows: int = 0
    semantics_indexed: int = 0
    docs_indexed: int = 0
    started_at: str = field(default_factory=_utcnow)
    finished_at: str = ""
    message: str = ""
    error: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "client_id": self.client_id,
            "status": self.status,
            "tables": self.tables,
            "rows_loaded": self.rows_loaded,
            "total_rows": self.total_rows,
            "semantics_indexed": self.semantics_indexed,
            "docs_indexed": self.docs_indexed,
            "started_at": self.started_at,
            "finished_at": self.finished_at,
            "message": self.message,
            "error": self.error,
        }


def _ch_client(host: str, port: int, password: str = "", database: str = "default"):
    import clickhouse_connect

    return clickhouse_connect.get_client(
        host=host, port=port, username="default", password=password, database=database
    )


def test_pg_connection(pg_dsn: str, schema: str = "public") -> dict[str, Any]:
    """Проверяет read-only доступ к Postgres клиента и возвращает список таблиц.

    Returns: {"ok": bool, "tables": [...], "error": str}
    """
    try:
        conn = connect_pg(pg_dsn)
        tables = list_tables(conn, schema)
        try:
            conn.close()
        except Exception:  # noqa: BLE001
            pass
        return {"ok": True, "tables": tables, "error": ""}
    except Exception as e:  # noqa: BLE001
        logger.warning("test_pg_connection: %s", e)
        return {"ok": False, "tables": [], "error": str(e)}


def run_tenant_etl(
    *,
    client_id: str,
    pg_dsn: str,
    pg_schema: str = "public",
    tables: list[str] | None = None,
    ch_host: str = "localhost",
    ch_port: int = 8123,
    ch_database: str | None = None,
    ch_password: str = "",
    add_client_id: bool = True,
    vector_collection: str | None = None,
    row_limit: int | None = None,
    progress: ProgressCb | None = None,
) -> EtlResult:
    """Полный ETL одного клиента: PG → ClickHouse → семантика.

    Все ресурсы (БД/таблицы/коллекция) создаются идемпотентно
    (``CREATE ... IF NOT EXISTS``). Перезаливка таблиц — full-refresh
    (TRUNCATE + INSERT), чтобы расписание давало консистентный слепок.
    """
    res = EtlResult(client_id=client_id)
    ch_database = ch_database or f"tenant_{client_id}"
    vector_collection = vector_collection or f"semantics_{client_id}"

    def emit(step: str, msg: str) -> None:
        logger.info("[%s] %s: %s", client_id, step, msg)
        if progress:
            try:
                progress(step, msg)
            except Exception:  # noqa: BLE001
                pass

    try:
        # 1. EXTRACT
        emit("extract", "подключение к Postgres клиента (read-only)")
        conn = connect_pg(pg_dsn)
        requested = [t.strip() for t in (tables or []) if t and t.strip()]
        table_list = requested or list_tables(conn, pg_schema)
        if not table_list:
            raise RuntimeError(f"В схеме '{pg_schema}' не найдено таблиц")
        res.tables = table_list
        emit("extract", f"таблиц к переносу: {len(table_list)} — {', '.join(table_list)}")

        schema_map: dict[str, list[dict]] = {}
        data_map: dict[str, tuple[list[str], list[tuple]]] = {}
        for tbl in table_list:
            schema_map[tbl] = get_columns(conn, pg_schema, tbl)
            data_map[tbl] = fetch_rows(conn, pg_schema, tbl, row_limit)
        try:
            conn.close()
        except Exception:  # noqa: BLE001
            pass

        # 2. LOAD
        emit("load", f"ClickHouse {ch_host}:{ch_port} → БД {ch_database}")
        ch_root = _ch_client(ch_host, ch_port, ch_password, "default")
        ch_root.command(f"CREATE DATABASE IF NOT EXISTS {ch_database}")
        ch_root.close()
        ch = _ch_client(ch_host, ch_port, ch_password, ch_database)

        ddls = {t: build_ch_ddl(t, cols, add_client_id) for t, cols in schema_map.items()}
        for tbl in table_list:
            ch.command(ddls[tbl])
            # full-refresh: чистим перед заливкой свежего слепка
            try:
                ch.command(f"TRUNCATE TABLE IF EXISTS {tbl}")
            except Exception as e:  # noqa: BLE001
                logger.debug("TRUNCATE %s пропущен: %s", tbl, e)
            cols, rows = data_map[tbl]
            if add_client_id:
                cols = ["client_id", *cols]
                rows = [(client_id, *r) for r in rows]
            if rows:
                ch.insert(tbl, [list(r) for r in rows], column_names=cols)
            res.rows_loaded[tbl] = len(rows)
            res.total_rows += len(rows)
            emit("load", f"таблица {tbl}: загружено {len(rows)} строк")

        # 3. SEMANTIC
        emit("semantic", "генерация семантического слоя по колонкам")
        semantics = build_semantics(schema_map, add_client_id)
        summary = index_semantics(semantics, vector_collection, ch)
        res.semantics_indexed = len(semantics)
        emit("semantic", summary)
        ch.close()

        res.status = "success"
        res.message = (
            f"{len(table_list)} табл., {res.total_rows} строк, "
            f"{res.semantics_indexed} семантических векторов"
        )
    except Exception as e:  # noqa: BLE001
        logger.exception("ETL клиента %s завершился ошибкой", client_id)
        res.status = "failed"
        res.error = str(e)
        res.message = f"Сбой ETL: {e}"
    finally:
        res.finished_at = _utcnow()
    return res


def ingest_tenant_docs(
    *,
    client_id: str,
    file_paths: list[str] | None = None,
    docs_dir: str | None = None,
    ch_host: str = "localhost",
    ch_port: int = 8123,
    ch_password: str = "",
    ch_database: str | None = None,
    collection: str | None = None,
    progress: ProgressCb | None = None,
) -> int:
    """Грузит документацию клиента в его RAG-коллекцию (таблица в ClickHouse клиента).

    Поддерживает PDF/TXT/MD/DOCX. Возвращает число проиндексированных чанков.
    Изоляция: документы пишутся в таблицу ``collection`` (по умолч. ``docs_<id>``)
    в персональной БД клиента, отдельно от общей базы знаний.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    from app.services.rag_service import get_embeddings_model

    ch_database = ch_database or f"tenant_{client_id}"
    collection = collection or f"docs_{client_id}"

    def emit(step: str, msg: str) -> None:
        logger.info("[%s] %s: %s", client_id, step, msg)
        if progress:
            try:
                progress(step, msg)
            except Exception:  # noqa: BLE001
                pass

    paths: list[str] = list(file_paths or [])
    if docs_dir and os.path.isdir(docs_dir):
        for fn in os.listdir(docs_dir):
            full = os.path.join(docs_dir, fn)
            if os.path.isfile(full):
                paths.append(full)
    if not paths:
        emit("docs", "нет файлов для загрузки")
        return 0

    documents = _load_documents(paths, emit)
    if not documents:
        emit("docs", "не удалось извлечь текст из документов")
        return 0

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = splitter.split_documents(documents)
    emit("docs", f"чанков получено: {len(splits)}")

    emb_model = get_embeddings_model()
    ch = _ch_client(ch_host, ch_port, ch_password, ch_database)
    ch.command(
        f"CREATE TABLE IF NOT EXISTS {collection} ("
        "id String, source String, content String, embedding Array(Float32)"
        ") ENGINE = MergeTree() ORDER BY id"
    )
    data = []
    for split in splits:
        emb = emb_model.embed_query(split.page_content)
        data.append(
            [
                uuid.uuid4().hex,
                os.path.basename(split.metadata.get("source", "doc")),
                split.page_content,
                emb,
            ]
        )
    if data:
        ch.insert(collection, data, column_names=["id", "source", "content", "embedding"])
    ch.close()
    emit("docs", f"проиндексировано {len(data)} чанков в '{collection}'")
    return len(data)


def _load_documents(paths: list[str], emit: ProgressCb):
    """Загружает PDF/TXT/MD/DOCX в список LangChain Document."""
    from langchain_core.documents import Document

    docs = []
    for path in paths:
        low = path.lower()
        try:
            if low.endswith(".pdf"):
                try:
                    import pdfplumber

                    with pdfplumber.open(path) as pdf:
                        for i, page in enumerate(pdf.pages):
                            text = (page.extract_text() or "").strip()
                            if text:
                                docs.append(
                                    Document(
                                        page_content=text,
                                        metadata={"source": path, "page": i},
                                    )
                                )
                except ImportError:
                    from langchain_community.document_loaders import PyPDFLoader

                    docs.extend(PyPDFLoader(path).load())
            elif low.endswith(".docx"):
                from docx import Document as DocxDocument

                d = DocxDocument(path)
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                if text:
                    docs.append(Document(page_content=text, metadata={"source": path}))
            else:  # txt / md / прочий текст
                from langchain_community.document_loaders import TextLoader

                docs.extend(TextLoader(path, encoding="utf-8").load())
        except Exception as e:  # noqa: BLE001
            emit("docs", f"пропущен {os.path.basename(path)}: {e}")
    return docs
